import docx
from docx.enum.text import WD_BREAK

doc = docx.Document("invitation.docx")

names = open("guests.txt")
names = names.read()
names = names.split(sep="\n")

for name in names:
    doc.add_paragraph("It would be a pleasure to have the company of").style = "inviteStyleLine1"
    doc.add_paragraph("RoboCop").style = "Name"
    doc.add_paragraph("at 11010 Memory Lane on the Evening of").style = "inviteStyleLine2"
    doc.add_paragraph("April 1st").style = "invitedate"
    date_line = doc.add_paragraph("at 7 o'clock")
    date_line.runs[0].add_break(break_type=WD_BREAK.PAGE)
    date_line.style = "inviteStyleLine2"

doc.save("invitations.docx")
